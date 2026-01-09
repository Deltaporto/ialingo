# -*- coding: utf-8 -*-
"""
Script para gerar DOCX do Plano de Trabalho CJF com formatação institucional.
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_cell_shading(cell, color):
    """Define cor de fundo de uma célula."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_horizontal_line(doc):
    """Adiciona linha horizontal."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("─" * 80)
    run.font.color.rgb = RGBColor(180, 180, 180)
    run.font.size = Pt(8)

def create_docx():
    doc = Document()

    # Configurar margens
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(3)

    # ===== TÍTULO =====
    title = doc.add_heading('Plano de Trabalho - Programa Equilibra (CJF)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ===== CABEÇALHO =====
    p = doc.add_paragraph()
    p.add_run('Unidade: ').bold = True
    p.add_run('3ª Relatoria da 2ª Turma Recursal - SJRJ\n')
    p.add_run('Gestor: ').bold = True
    p.add_run('Dr. Rafael Assis Alves (em exercício de titularidade)\n')
    p.add_run('Data: ').bold = True
    p.add_run('Janeiro/2026')

    add_horizontal_line(doc)

    # ===== BLOCO REQUISITÓRIO =====
    req_table = doc.add_table(rows=1, cols=1)
    req_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = req_table.rows[0].cells[0]
    set_cell_shading(cell, 'F0F0F0')

    p = cell.paragraphs[0]
    p.add_run('Documento Requisitório: ').bold = True
    p.add_run('Ofício TRF2 1452551, de 19/12/2025\n')
    p.add_run('Solicitantes: ').bold = True
    p.add_run('Des. Roy Reis Friede (Coordenador JEF) e Des. Wanderley Sanan (Vice)\n')
    p.add_run('Prazo de Entrega: ').bold = True
    p.add_run('13/01/2026').bold = True
    p.add_run('\n')
    p.add_run('Objetivo: ').bold = True
    p.add_run('Demonstrar a trajetória de recuperação da unidade, os resultados alcançados e as perspectivas para 2026.')

    doc.add_paragraph()

    # ===== SEÇÃO 1 =====
    doc.add_heading('1. Diagnóstico Situacional: Fatores Determinantes do Passivo (2024)', 1)

    p = doc.add_paragraph()
    p.add_run('A oscilação de produtividade verificada no exercício de 2024 não decorreu de incremento de distribuição ou ineficiência estrutural, mas de um ')
    p.add_run('evento crítico de descontinuidade administrativa e redução da força de trabalho').bold = True
    p.add_run('.')

    # 1.1
    doc.add_heading('1.1. Descontinuidade na Gestão e Liderança', 2)
    p = doc.add_paragraph('Documentos oficiais (Ofícios JFRJ 2023/04656 e 2023/05644) evidenciam um cenário de vacância simultânea na gestão da unidade:')

    items_11 = [
        ('Aposentadoria do Magistrado Titular', ' (Dr. Paulo Alberto Jorge, em fev/2024).'),
        ('Remoção da Chefia de Gabinete', ' (Anna Beatriz, 13 anos de experiência na unidade), transferida para a Secretaria Única em jan/2024.'),
        ('Evasão de Capital Intelectual', ': Ocorre a saída de servidores experientes para outras unidades, gerando perda de memória técnica:\n    • Analista Bruno Vivas (3 anos de lotação) removido para o Gab-01.\n    • Técnica Carina Pacheco (11 anos de lotação) removida para o Gab-02.'),
        ('Regime de Acumulação', ': No período crítico, a unidade foi presidida por magistrada designada em regime de acumulação de acervo ("sem prejuízo da origem"), o que inviabilizou a gestão exclusiva necessária para o momento de crise (Vide Ato TRF2-ATC-2024/00040).'),
        ('Desqualificação Técnica Momentânea', ': A equipe remanescente contava com apenas um servidor (Marcelo) apto para a elaboração de minutas.'),
    ]

    for i, (bold_part, rest) in enumerate(items_11, 1):
        p = doc.add_paragraph(style='List Number')
        p.add_run(bold_part).bold = True
        p.add_run(rest)

    # 1.2
    doc.add_heading('1.2. Impacto Quantitativo nos Indicadores', 2)
    p = doc.add_paragraph('O reflexo operacional foi imediato:')

    items_12 = [
        ('Demanda Constante', ': Distribuição estável (1.823 em 2023 → 1.854 em 2024).'),
        ('Retração Produtiva', ': Queda de 58% no volume de julgamentos (1.635 para 684), decorrente da ausência de pessoal.'),
        ('Formação de Passivo', ': O acervo acumulado foi consequência direta da redução da capacidade de processamento (output).'),
    ]

    for bold_part, rest in items_12:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(bold_part).bold = True
        p.add_run(rest)

    # 1.3
    doc.add_heading('1.3. Distorção Qualitativa do Acervo Herdado', 2)
    p = doc.add_paragraph()
    p.add_run('A análise temática revela que a equipe remanescente de 2024, operando em modo de sobrevivência, priorizou matérias de menor complexidade (BPC/LOAS). Consequentemente, a gestão atual herdou um acervo com concentração anormal de casos de alta complexidade previdenciária (aposentadoria especial, tempo especial) — um passivo não apenas quantitativo, mas ')
    p.add_run('qualitativamente mais difícil').bold = True
    p.add_run(' que a média histórica da unidade.')

    add_horizontal_line(doc)

    # ===== SEÇÃO 2 =====
    doc.add_heading('2. Plano de Reestruturação e Gestão Estratégica (2025)', 1)

    p = doc.add_paragraph()
    p.add_run('Em ')
    p.add_run('Setembro/2024').bold = True
    p.add_run(', a atual gestão iniciou o processo de reconstrução das rotinas de trabalho.')

    # 2.1
    doc.add_heading('2.1. Metodologia Aplicada', 2)

    items_21 = [
        ('Priorização Racional ("70/30")', ': Alocação majoritária da força de trabalho (70%) para redução do acervo com paralisação > 150 dias.'),
        ('Gestão Visual de Fluxo', ': Implementação de localizadores no sistema Eproc para controle de prazos.'),
        ('Capacitação Técnica Intensiva', ': Assunção direta, pelo magistrado, do treinamento dos novos servidores e estagiários.'),
    ]

    for bold_part, rest in items_21:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(bold_part).bold = True
        p.add_run(rest)

    # 2.2
    doc.add_heading('2.2. Resultados Alcançados: Da Inclusão no PAE à Liderança Regional', 2)
    p = doc.add_paragraph('A eficácia das medidas é comprovada pelos indicadores oficiais do TRF2:')

    # BOX DE DESTAQUE
    doc.add_paragraph()
    box_table = doc.add_table(rows=1, cols=1)
    box_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    box_cell = box_table.rows[0].cells[0]
    set_cell_shading(box_cell, '1F4E79')  # Azul escuro

    p = box_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('RESULTADO INSTITUCIONAL\n\n')
    run.bold = True
    run.font.color.rgb = RGBColor(255, 255, 255)
    run.font.size = Pt(14)

    run = p.add_run('1º LUGAR no cumprimento da Meta 1 do CNJ\n')
    run.bold = True
    run.font.color.rgb = RGBColor(255, 215, 0)  # Dourado
    run.font.size = Pt(16)

    run = p.add_run('entre todas as 24 Relatorias das Turmas Recursais do RJ\n\n')
    run.font.color.rgb = RGBColor(255, 255, 255)
    run.font.size = Pt(11)

    run = p.add_run('Índice: 131,46% | Fonte: Quadro de Acompanhamento PAE (Jan/2026)')
    run.font.color.rgb = RGBColor(200, 200, 200)
    run.font.size = Pt(10)

    doc.add_paragraph()

    items_22 = [
        ('Saneamento do Passivo', ': Estoque crítico (parados > 150 dias) zerado em Julho/2025 e mantido sob controle.'),
        ('Redução Global', ': Diminuição de 35% no acervo total (de 2.056 para 1.329).'),
    ]

    for bold_part, rest in items_22:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(bold_part).bold = True
        p.add_run(rest)

    p = doc.add_paragraph()
    p.add_run('Em menos de 16 meses, a unidade transitou da condição de "unidade sob acompanhamento especial" para ')
    p.add_run('referência de eficiência').bold = True
    p.add_run(' no sistema de Turmas Recursais do Rio de Janeiro.')

    # 2.3
    doc.add_heading('2.3. Comprovação Visual: A Trajetória nos Indicadores Oficiais', 2)
    p = doc.add_paragraph('Os gráficos extraídos do Painel CNJ (Justiça em Números) ilustram a trajetória da unidade de forma inequívoca:')

    # Tabela de indicadores
    table = doc.add_table(rows=3, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['Indicador', 'Pior Momento (2024)', 'Situação Atual (Nov/2025)', 'Variação']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'D9E2F3')

    data = [
        ['Índice de Atendimento à Demanda', '41,26%', '113,92%', '+176%'],
        ['Taxa de Congestionamento Líquida', '69,76%', '33,86%', '-51%'],
    ]

    for row_idx, row_data in enumerate(data, 1):
        for col_idx, value in enumerate(row_data):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = value
            if col_idx == 0:
                cell.paragraphs[0].runs[0].bold = True

    p = doc.add_paragraph()
    p.add_run('Gráficos anexos: CNJ_IAD.png e CNJ_TCL.png').italic = True

    add_horizontal_line(doc)

    # ===== SEÇÃO 3 =====
    doc.add_heading('3. Perspectivas e Metas de Convergência (2026)', 1)

    p = doc.add_paragraph('O objetivo agora deixa de ser "apagar incêndio" e passa a ser "manter a velocidade de cruzeiro", consolidando os resultados alcançados.')

    # 3.1
    doc.add_heading('3.1. Consolidação da Equipe e Continuidade Operacional', 2)
    p = doc.add_paragraph('Os servidores do Grupo de Apoio (GSA), cuja prorrogação foi deferida, constituem hoje recurso humano central para a manutenção dos resultados:')

    items_31 = [
        ('Participação na Produção', ': Responsáveis por 28,2% de todos os atos decisórios de 2025.'),
        ('Qualificação Consolidada', ': Após 12 meses de capacitação, dominam a instrução de demandas de alta complexidade previdenciária.'),
    ]

    for bold_part, rest in items_31:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(bold_part).bold = True
        p.add_run(rest)

    p = doc.add_paragraph('A estabilidade da equipe atual assegura a sustentabilidade da produtividade demonstrada.')

    # 3.2
    doc.add_heading('3.2. Metas de Convergência (Ofício CJF 0805981)', 2)
    p = doc.add_paragraph('Para atingir a média da 2ª Região, o plano estabelece os seguintes alvos numéricos para 2026:')

    # Tabela de metas
    table = doc.add_table(rows=4, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['Indicador', 'Situação Atual (Out/25)', 'Meta 2026 (Alvo)', 'Ação Necessária']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'D9E2F3')

    data = [
        ['Acervo Total', '1.487', '~733 (Média RJ)', 'Reduzir 50% do estoque via mutirão GSA.'],
        ['Tempo Médio', '18,0 meses', '~15,7 meses', 'Priorizar baixas antigas (Meta 2).'],
        ['Sem 1ª Decisão', '1.048', '~465', 'Manter ritmo de 300+ sentenças/mês.'],
    ]

    for row_idx, row_data in enumerate(data, 1):
        for col_idx, value in enumerate(row_data):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = value
            if col_idx == 0:
                cell.paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    # Bloco de estratégia
    strat_table = doc.add_table(rows=1, cols=1)
    strat_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = strat_table.rows[0].cells[0]
    set_cell_shading(cell, 'E2EFDA')  # Verde claro

    p = cell.paragraphs[0]
    p.add_run('Estratégia: ').bold = True
    p.add_run('A manutenção da produtividade atual (3.250 atos/ano) permitirá atingir a média regional em ')
    p.add_run('aproximadamente 8 meses').bold = True
    p.add_run(', desde que a força de trabalho seja mantida integralmente (incluindo GSA).')

    # 3.3
    doc.add_heading('3.3. Condições para Manutenção da Trajetória', 2)
    p = doc.add_paragraph('A continuidade dos resultados depende de:')

    items_33 = [
        ('Estabilidade do quadro atual', ': Manutenção dos servidores GSA (já prorrogados) durante o período de maturação dos novos estagiários.'),
        ('Cobertura de afastamentos', ': Substituição tempestiva durante licenças programadas (evitando oscilações de produtividade como a verificada em Nov/2025).'),
    ]

    for i, (bold_part, rest) in enumerate(items_33, 1):
        p = doc.add_paragraph(style='List Number')
        p.add_run(bold_part).bold = True
        p.add_run(rest)

    add_horizontal_line(doc)

    # ===== CONCLUSÃO =====
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run('A trajetória documentada demonstra que a crise de 2024 decorreu de fatores circunstanciais de pessoal — não de falhas estruturais ou de gestão. A recuperação de 2025, coroada pelo 1º lugar no ranking regional de eficiência, comprova a eficácia da metodologia aplicada e o retorno do investimento em capacitação. A unidade encontra-se, hoje, em condições de atingir a média regional no exercício de 2026.')
    run.italic = True

    # Salvar
    output_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'docs', 'gabinete', 'PLANO_TRABALHO_CJF.docx')
    doc.save(output_path)
    print(f"DOCX gerado com sucesso: {output_path}")
    return output_path

if __name__ == '__main__':
    create_docx()

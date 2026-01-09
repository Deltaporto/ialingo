# -*- coding: utf-8 -*-
"""
Script para gerar DOCX do Plano de Trabalho CJF em formato de prosa institucional.
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_cell_shading(cell, color):
    """Define cor de fundo de uma célula."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_horizontal_line(doc):
    """Adiciona linha horizontal."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("─" * 80)
    run.font.color.rgb = RGBColor(180, 180, 180)
    run.font.size = Pt(8)

def add_justified_paragraph(doc, text, space_after=Pt(12)):
    """Adiciona parágrafo justificado."""
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = space_after
    p.paragraph_format.line_spacing = 1.15
    return p

def create_docx():
    doc = Document()

    # Configurar margens
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(3)

    # ===== TÍTULO =====
    title = doc.add_heading('Plano de Trabalho - Programa Equilibra (CJF)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ===== CABEÇALHO =====
    p = doc.add_paragraph()
    p.add_run('Unidade: ').bold = True
    p.add_run('3ª Relatoria da 2ª Turma Recursal - SJRJ\n')
    p.add_run('Gestor: ').bold = True
    p.add_run('Dr. Rafael Assis Alves (em exercício de titularidade)\n')
    p.add_run('Data: ').bold = True
    p.add_run('Janeiro/2026')

    add_horizontal_line(doc)

    # ===== BLOCO REQUISITÓRIO =====
    req_table = doc.add_table(rows=1, cols=1)
    req_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = req_table.rows[0].cells[0]
    set_cell_shading(cell, 'F0F0F0')

    p = cell.paragraphs[0]
    p.add_run('Documento Requisitório: ').bold = True
    p.add_run('Ofício TRF2 1452551, de 19/12/2025\n')
    p.add_run('Solicitantes: ').bold = True
    p.add_run('Des. Roy Reis Friede (Coordenador JEF) e Des. Wanderley Sanan (Vice)\n')
    p.add_run('Prazo de Entrega: ').bold = True
    p.add_run('13/01/2026').bold = True
    p.add_run('\n')
    p.add_run('Objetivo: ').bold = True
    p.add_run('Demonstrar a trajetória de recuperação da unidade, os resultados alcançados e as perspectivas para 2026.')

    doc.add_paragraph()

    # ===== SEÇÃO 1 =====
    doc.add_heading('1. Diagnóstico Situacional: Fatores Determinantes do Passivo (2024)', 1)

    add_justified_paragraph(doc, 'A oscilação de produtividade verificada no exercício de 2024 não decorreu de incremento de distribuição ou ineficiência estrutural, mas de um evento crítico de descontinuidade administrativa conjugado com expressiva redução da força de trabalho qualificada. A análise dos indicadores de desempenho da unidade, quando cotejada com a documentação administrativa oficial, evidencia um cenário de ruptura operacional cujas consequências repercutiram de forma significativa na capacidade de processamento jurisdicional.')

    # 1.1
    doc.add_heading('1.1. Descontinuidade na Gestão e Liderança', 2)

    add_justified_paragraph(doc, 'A descontinuidade na gestão da 3ª Relatoria decorreu de múltiplos fatores concomitantes, todos documentados em ofícios oficiais (JFRJ 2023/04656 e 2023/05644), que configuraram um cenário de vacância simultânea nos quadros de liderança e execução técnica. A aposentadoria do magistrado titular, Dr. Paulo Alberto Jorge, efetivada em fevereiro de 2024, somou-se à transferência da Chefe de Gabinete, Anna Beatriz Silva de Lima Pereira — profissional com treze anos de experiência na unidade —, removida para a Secretaria Única em janeiro do mesmo ano. Essa perda na camada de gestão foi agravada por um fenômeno paralelo de evasão de capital intelectual: a remoção de servidores experientes para outras unidades gerou significativa perda de memória técnica institucional. O Analista Judiciário Bruno Vivas, com três anos de lotação no Gabinete 3, foi removido para o Gabinete 1, enquanto a Técnica Judiciária Carina Pacheco, com onze anos de lotação na unidade, foi transferida para o Gabinete 2.')

    add_justified_paragraph(doc, 'Paralelamente a essa erosão de quadros, o período crítico de transição foi marcado por regime de gestão precário: a unidade foi presidida por magistrada designada em regime de acumulação de acervo, operando "sem prejuízo da origem" (Ato TRF2-ATC-2024/00040), modelo administrativo que inviabilizou a gestão exclusiva e intensiva demandada pelo momento de crise. Ademais, a desqualificação técnica momentânea da equipe remanescente tornou-se evidente: dos servidores mantidos na lotação, apenas um — o servidor Marcelo — encontrava-se apto para a elaboração de minutas de decisão, comprometendo severamente a capacidade produtiva da unidade naquele período crítico de transição.')

    # 1.2
    doc.add_heading('1.2. Impacto Quantitativo nos Indicadores', 2)

    add_justified_paragraph(doc, 'O reflexo operacional dessas rupturas administrativas manifestou-se de forma imediata nos indicadores de desempenho da unidade. A demanda jurisdicional manteve-se estável, com distribuição de 1.823 processos em 2023 e 1.854 em 2024, o que afasta qualquer hipótese de incremento anormal de carga de trabalho como fator explicativo da crise. Entretanto, verificou-se retração produtiva acentuada: o volume de julgamentos sofreu queda de 58% em termos absolutos, passando de 1.635 decisões em 2023 para apenas 684 em 2024. Essa redução drástica no output decisório decorreu diretamente da ausência de pessoal qualificado para instrução processual e elaboração de minutas.')

    add_justified_paragraph(doc, 'A formação do passivo processual, portanto, configurou-se como consequência inexorável da redução da capacidade de processamento, uma vez que a entrada de processos (input) permaneceu constante enquanto a saída (output) foi drasticamente comprimida pela insuficiência de força de trabalho. O acervo acumulado representa, assim, não um problema de gestão inadequada, mas o resultado matemático da disparidade entre demanda constante e capacidade operacional temporariamente reduzida.')

    # 1.3
    doc.add_heading('1.3. Distorção Qualitativa do Acervo Herdado', 2)

    add_justified_paragraph(doc, 'A análise temática do acervo processual da unidade revela que os efeitos da crise de 2024 transcenderam a dimensão quantitativa, produzindo também uma distorção qualitativa significativa. A equipe remanescente, operando em regime de sobrevivência administrativa e com capacidade técnica limitada, adotou estratégia de priorização baseada na complexidade das matérias: concentrou-se na instrução e julgamento de demandas de menor complexidade jurídica, notadamente processos relativos ao Benefício de Prestação Continuada (BPC/LOAS), que demandam menor tempo de análise e cujo processamento pode ser conduzido com menor expertise técnica.')

    add_justified_paragraph(doc, 'Consequentemente, a atual gestão herdou um acervo com concentração anormal de casos de alta complexidade previdenciária, especialmente processos envolvendo reconhecimento de tempo especial e concessão de aposentadoria especial, matérias que exigem análise técnica aprofundada de laudos periciais, legislação previdenciária específica e jurisprudência especializada. Trata-se, portanto, de um passivo não apenas quantitativamente expressivo, mas qualitativamente mais difícil que a média histórica da unidade, demandando investimento de tempo significativamente superior por processo para saneamento adequado. Essa distorção qualitativa representa desafio adicional ao processo de recuperação da unidade, uma vez que a redução do acervo não pode ser alcançada mediante mero incremento quantitativo de decisões, mas exige análise técnica aprofundada e capacitação específica da equipe para enfrentamento de matérias de maior densidade jurídica.')

    add_horizontal_line(doc)

    # ===== SEÇÃO 2 =====
    doc.add_heading('2. Plano de Reestruturação e Gestão Estratégica (2025)', 1)

    add_justified_paragraph(doc, 'Em setembro de 2024, a atual gestão iniciou o processo de reconstrução integral das rotinas de trabalho da unidade. O cenário herdado exigia não apenas a superação de déficit quantitativo, mas a reestruturação completa da capacidade operacional do gabinete, comprometida pela descontinuidade administrativa descrita na seção precedente.')

    # 2.1
    doc.add_heading('2.1. Metodologia Aplicada', 2)

    add_justified_paragraph(doc, 'A estratégia de recuperação fundamentou-se em três pilares metodológicos complementares, concebidos para maximizar a eficiência da equipe reduzida e reverter os indicadores críticos no menor prazo possível. Adotou-se, inicialmente, o modelo de priorização racional baseado na regra "70/30", mediante o qual 70% da força de trabalho disponível foi alocada especificamente para a redução do acervo com tempo de paralisação superior a 150 dias — justamente o contingente que ensejara a inclusão da unidade no Plano de Ação Especial (PAE). Esta concentração de esforços permitiu o saneamento acelerado do passivo crítico, ao mesmo tempo em que os 30% remanescentes da capacidade produtiva eram direcionados ao atendimento do fluxo ordinário de distribuição, evitando a formação de novo estoque de casos paralisados.')

    add_justified_paragraph(doc, 'Paralelamente, implementou-se sistema de gestão visual de fluxo mediante a criação e utilização intensiva de localizadores no sistema Eproc. Essa ferramenta de controle gerencial permitiu o monitoramento em tempo real dos prazos processuais, a identificação imediata de gargalos operacionais e a redistribuição dinâmica de carga de trabalho entre os integrantes da equipe. A terceira vertente da metodologia — talvez a mais determinante para o êxito do plano — consistiu na assunção direta, pelo magistrado titular, da capacitação técnica intensiva dos novos servidores e estagiários. Diante da ausência de quadro experiente capaz de transmitir a memória técnica da unidade, o magistrado conduziu pessoalmente o treinamento em instrução processual e elaboração de minutas decisórias, acelerando substancialmente a curva de aprendizado da equipe recém-constituída.')

    # 2.2
    doc.add_heading('2.2. Resultados Alcançados: Da Inclusão no PAE à Liderança Regional', 2)

    add_justified_paragraph(doc, 'A eficácia das medidas implementadas pode ser objetivamente aferida pelos indicadores oficiais consolidados pelo Tribunal Regional Federal da 2ª Região. A trajetória percorrida pela unidade no período de setembro de 2024 a dezembro de 2025 demonstra não apenas a recuperação dos patamares produtivos anteriores, mas a superação das médias regionais em diversos indicadores estratégicos de desempenho.')

    # BOX DE DESTAQUE
    doc.add_paragraph()
    box_table = doc.add_table(rows=1, cols=1)
    box_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    box_cell = box_table.rows[0].cells[0]
    set_cell_shading(box_cell, '1F4E79')

    p = box_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('RESULTADO INSTITUCIONAL\n\n')
    run.bold = True
    run.font.color.rgb = RGBColor(255, 255, 255)
    run.font.size = Pt(14)

    run = p.add_run('1º LUGAR no cumprimento da Meta 1 do CNJ\n')
    run.bold = True
    run.font.color.rgb = RGBColor(255, 215, 0)
    run.font.size = Pt(16)

    run = p.add_run('entre todas as 24 Relatorias das Turmas Recursais do RJ\n\n')
    run.font.color.rgb = RGBColor(255, 255, 255)
    run.font.size = Pt(11)

    run = p.add_run('Índice: 131,46% | Fonte: Quadro de Acompanhamento PAE (Jan/2026)')
    run.font.color.rgb = RGBColor(200, 200, 200)
    run.font.size = Pt(10)

    doc.add_paragraph()

    add_justified_paragraph(doc, 'O saneamento do passivo crítico foi concluído em julho de 2025, quando o estoque de processos paralisados há mais de 150 dias foi integralmente eliminado, sendo mantido sob controle rigoroso desde então. Concomitantemente, verificou-se redução de 35% no acervo total da unidade, que regrediu de 2.056 processos para 1.329 processos no período compreendido entre setembro de 2024 e novembro de 2025.')

    add_justified_paragraph(doc, 'Em menos de dezesseis meses, a unidade transitou da condição de "unidade sob acompanhamento especial", incluída no PAE em razão de indicadores críticos de congestionamento, para a posição de referência de eficiência no sistema de Turmas Recursais do Estado do Rio de Janeiro. A conquista do primeiro lugar no cumprimento da Meta 1 do Conselho Nacional de Justiça, com índice de 131,46% entre as vinte e quatro relatorias do Estado, materializa institucionalmente o êxito da estratégia de recuperação implementada.')

    # 2.3
    doc.add_heading('2.3. Comprovação Visual: A Trajetória nos Indicadores Oficiais', 2)

    add_justified_paragraph(doc, 'Os dados extraídos do Painel CNJ (Justiça em Números) e dos relatórios gerenciais do Tribunal Regional Federal ilustram, de forma inequívoca, a magnitude da transformação operacional vivenciada pela unidade no período compreendido entre o segundo semestre de 2024 e o final de 2025.')

    # Tabela de indicadores
    table = doc.add_table(rows=3, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['Indicador', 'Pior Momento (2024)', 'Situação Atual (Nov/2025)', 'Variação']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'D9E2F3')

    data = [
        ['Índice de Atendimento à Demanda', '41,26%', '113,92%', '+176%'],
        ['Taxa de Congestionamento Líquida', '69,76%', '33,86%', '-51%'],
    ]

    for row_idx, row_data in enumerate(data, 1):
        for col_idx, value in enumerate(row_data):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = value
            if col_idx == 0:
                cell.paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    add_justified_paragraph(doc, 'O Índice de Atendimento à Demanda (IAD), que mensura a relação entre processos baixados e casos novos distribuídos, evidencia a reversão completa do quadro de insuficiência produtiva. A evolução de 41,26% para 113,92% representa incremento de 176 pontos percentuais, indicando que a unidade não apenas equiparou sua capacidade de processamento à demanda de distribuição, mas passou a operar em ritmo superior ao necessário para estabilização do acervo, permitindo a redução ativa do estoque acumulado.')

    add_justified_paragraph(doc, 'Paralelamente, a Taxa de Congestionamento Líquida (TCL) — indicador que mede o percentual de processos que permaneceram em tramitação sem solução definitiva ao longo do exercício — experimentou redução de 51 pontos percentuais, transitando de 69,76% para 33,86%. Esta redução situa a unidade em patamar inferior à média nacional das Turmas Recursais (que se mantém historicamente acima de 50%), demonstrando não apenas a recuperação, mas a excelência do desempenho atual quando comparado aos padrões nacionais de produtividade judiciária.')

    add_justified_paragraph(doc, 'A documentação fotográfica dos painéis oficiais do CNJ, anexa a este plano, permite a verificação visual da trajetória ascendente dos indicadores ao longo dos meses de 2025, evidenciando a consistência e sustentabilidade dos resultados alcançados.')

    p = doc.add_paragraph()
    p.add_run('Gráficos anexos: CNJ_IAD.png e CNJ_TCL.png').italic = True

    add_horizontal_line(doc)

    # ===== SEÇÃO 3 =====
    doc.add_heading('3. Perspectivas e Metas de Convergência (2026)', 1)

    add_justified_paragraph(doc, 'O objetivo estratégico para o exercício de 2026 deixa de ser a recuperação emergencial do acervo crítico e passa a concentrar-se na consolidação dos resultados alcançados, assegurando a manutenção da velocidade de cruzeiro demonstrada no último exercício. Trata-se, agora, de transformar os ganhos quantitativos e qualitativos obtidos em 2025 em capacidade permanente de resposta jurisdicional, garantindo a convergência gradual dos indicadores da unidade em direção à média regional.')

    # 3.1
    doc.add_heading('3.1. Consolidação da Equipe e Continuidade Operacional', 2)

    add_justified_paragraph(doc, 'Os servidores do Grupo de Apoio (GSA), cuja prorrogação foi deferida pela administração superior, constituem hoje recurso humano central para a manutenção dos resultados demonstrados ao longo de 2025. Esses profissionais responderam por 28,2% de todos os atos decisórios produzidos pela unidade no exercício anterior, representando participação expressiva na capacidade produtiva instalada. Após 12 meses de capacitação técnica intensiva, conduzida diretamente pelo magistrado titular, a equipe atual domina plenamente a instrução de demandas de alta complexidade previdenciária, incluindo matérias que demandam análise aprofundada de legislação especializada, jurisprudência superior e produção probatória técnica. A estabilidade desse núcleo qualificado assegura a sustentabilidade da produtividade demonstrada e permite a condução simultânea do processo de formação dos novos estagiários, sem comprometimento dos níveis atuais de atendimento à demanda.')

    # 3.2
    doc.add_heading('3.2. Metas de Convergência (Ofício CJF 0805981)', 2)

    add_justified_paragraph(doc, 'Para atingir a média da 2ª Região nos principais indicadores de desempenho monitorados pelo Conselho da Justiça Federal, o plano de trabalho estabelece os seguintes alvos numéricos para o exercício de 2026:')

    # Tabela de metas
    table = doc.add_table(rows=4, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['Indicador', 'Situação Atual (Out/25)', 'Meta 2026 (Alvo)', 'Ação Necessária']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'D9E2F3')

    data = [
        ['Acervo Total', '1.487', '~733 (Média RJ)', 'Reduzir 50% do estoque via mutirão GSA.'],
        ['Tempo Médio', '18,0 meses', '~15,7 meses', 'Priorizar baixas antigas (Meta 2).'],
        ['Sem 1ª Decisão', '1.048', '~465', 'Manter ritmo de 300+ sentenças/mês.'],
    ]

    for row_idx, row_data in enumerate(data, 1):
        for col_idx, value in enumerate(row_data):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = value
            if col_idx == 0:
                cell.paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    # Bloco de estratégia
    strat_table = doc.add_table(rows=1, cols=1)
    strat_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = strat_table.rows[0].cells[0]
    set_cell_shading(cell, 'E2EFDA')

    p = cell.paragraphs[0]
    p.add_run('Estratégia: ').bold = True
    p.add_run('A manutenção da produtividade atual (3.250 atos/ano) permitirá atingir a média regional em ')
    p.add_run('aproximadamente 8 meses').bold = True
    p.add_run(', desde que a força de trabalho seja mantida integralmente (incluindo GSA).')

    # 3.3
    doc.add_heading('3.3. Condições para Manutenção da Trajetória', 2)

    add_justified_paragraph(doc, 'A continuidade dos resultados obtidos em 2025 e a efetiva convergência aos patamares médios regionais no exercício de 2026 dependem, fundamentalmente, da estabilidade do quadro atual de pessoal durante o período de maturação técnica dos novos estagiários recentemente incorporados à unidade. A manutenção dos servidores do Grupo de Apoio (GSA), cuja permanência já foi prorrogada pela administração, constitui condição operacional indispensável para assegurar que a capacidade produtiva instalada não sofra descontinuidade abrupta antes que a nova equipe alcance o grau de qualificação necessário para assumir, de forma autônoma, a instrução de demandas complexas. Complementarmente, a cobertura tempestiva de eventuais afastamentos programados — mediante substituição ou redistribuição de tarefas — mostra-se igualmente necessária para evitar oscilações de produtividade como a verificada em novembro de 2025, quando a ausência momentânea de servidor-chave impactou, ainda que transitoriamente, o ritmo de julgamentos. A manutenção dessas condições estruturais permitirá que a unidade atinja, de forma sustentável e progressiva, a média regional nos indicadores monitorados pelo Conselho da Justiça Federal, consolidando a trajetória de recuperação institucional documentada neste plano.')

    add_horizontal_line(doc)

    # ===== CONCLUSÃO =====
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run('A trajetória documentada demonstra que a crise de 2024 decorreu de fatores circunstanciais de pessoal — não de falhas estruturais ou de gestão. A recuperação de 2025, coroada pelo 1º lugar no ranking regional de eficiência, comprova a eficácia da metodologia aplicada e o retorno do investimento em capacitação. A unidade encontra-se, hoje, em condições de atingir a média regional no exercício de 2026.')
    run.italic = True

    # Salvar
    output_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'docs', 'gabinete', 'PLANO_TRABALHO_CJF_PROSA.docx')
    doc.save(output_path)
    print(f"DOCX gerado com sucesso: {output_path}")
    return output_path

if __name__ == '__main__':
    create_docx()
